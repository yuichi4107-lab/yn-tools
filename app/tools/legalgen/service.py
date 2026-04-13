"""契約書・利用規約自動作成ツール - AI生成ロジック + Word/PDF出力"""

import io
import json
import re
import tempfile
from pathlib import Path

from openai import AsyncOpenAI

from app.config import settings

_client: AsyncOpenAI | None = None

DISCLAIMER = (
    "※ 本文書はAIによって自動生成されたものであり、法的助言ではありません。"
    "実際の法的効力を担保するには、弁護士等の専門家にご確認ください。"
)


def _get_client() -> AsyncOpenAI:
    global _client
    if _client is None:
        _client = AsyncOpenAI(api_key=settings.openai_api_key)
    return _client


# 文書種別定義
DOC_TYPES: dict[str, dict] = {
    "commission": {
        "name": "業務委託契約書",
        "icon": "📋",
        "color": "blue",
        "required_clauses": [
            "業務の内容と範囲",
            "委託料（報酬）と支払条件",
            "業務遂行期間",
            "知的財産権の帰属",
            "秘密保持義務",
            "損害賠償",
            "契約の解除",
            "準拠法と合意管轄",
        ],
        "specific_fields": ["work_content", "compensation", "period", "ip_ownership"],
    },
    "nda": {
        "name": "NDA（秘密保持契約書）",
        "icon": "🔒",
        "color": "purple",
        "required_clauses": [
            "秘密情報の定義と範囲",
            "秘密保持義務の内容",
            "秘密情報の使用目的",
            "秘密保持期間",
            "秘密情報の返還・廃棄",
            "例外事項（公知情報等）",
            "損害賠償",
            "準拠法と合意管轄",
        ],
        "specific_fields": ["purpose", "info_scope", "period"],
    },
    "sale": {
        "name": "売買契約書",
        "icon": "🛒",
        "color": "green",
        "required_clauses": [
            "売買の目的物",
            "売買代金と支払方法",
            "引渡条件",
            "所有権移転",
            "危険負担",
            "瑕疵担保責任（契約不適合責任）",
            "解除条件",
            "準拠法と合意管轄",
        ],
        "specific_fields": ["product_name", "amount", "delivery_conditions"],
    },
    "tos": {
        "name": "利用規約",
        "icon": "📜",
        "color": "orange",
        "required_clauses": [
            "サービスの目的と提供内容",
            "利用資格と登録",
            "禁止事項",
            "免責事項",
            "知的財産権",
            "個人情報の取扱い",
            "サービスの変更・停止",
            "準拠法と管轄",
        ],
        "specific_fields": ["service_name", "prohibited_items", "disclaimer"],
    },
    "privacy": {
        "name": "プライバシーポリシー",
        "icon": "🛡️",
        "color": "teal",
        "required_clauses": [
            "取得する個人情報の項目",
            "個人情報の利用目的",
            "個人情報の第三者提供",
            "個人情報の管理",
            "Cookieの使用",
            "開示・訂正・削除の請求",
            "問い合わせ先",
            "ポリシーの改定",
        ],
        "specific_fields": ["collected_info", "usage_purpose", "contact_info"],
    },
    "employment": {
        "name": "雇用契約書",
        "icon": "👔",
        "color": "indigo",
        "required_clauses": [
            "雇用期間（有期・無期）",
            "業務内容",
            "勤務場所",
            "勤務時間・休日",
            "賃金と支払方法",
            "社会保険・福利厚生",
            "解雇・退職の手続き",
            "準拠法と合意管轄",
        ],
        "specific_fields": ["work_content", "salary", "work_location", "period"],
    },
    "rent": {
        "name": "賃貸借契約書",
        "icon": "🏠",
        "color": "red",
        "required_clauses": [
            "賃貸物件の表示",
            "賃料と支払方法",
            "敷金・保証金",
            "契約期間",
            "使用目的",
            "修繕・費用負担",
            "禁止事項",
            "解約・明渡条件",
            "準拠法と合意管轄",
        ],
        "specific_fields": ["property", "rent", "period", "deposit"],
    },
}

# 文書種別ごとのシステムプロンプト
SYSTEM_PROMPTS: dict[str, str] = {
    "commission": """\
あなたは日本の契約実務に精通した法律文書作成の専門家です。
以下の条件で業務委託契約書を作成してください。

【必ず含める条項】
1. 業務の内容と範囲
2. 委託料（報酬）と支払条件
3. 業務遂行期間
4. 知的財産権の帰属
5. 秘密保持義務
6. 損害賠償
7. 契約の解除
8. 準拠法と合意管轄

【出力形式】
- Markdown形式で出力（## 第X条 条項名、本文）
- 条項番号は「第1条」「第2条」...と順番に付ける
- 正式な法律文書の文体（甲・乙を使用）
- AIによる断定的な法的判断（「この契約は有効です」等）は記載しない
""",
    "nda": """\
あなたは日本の契約実務に精通した法律文書作成の専門家です。
以下の条件でNDA（秘密保持契約書）を作成してください。

【必ず含める条項】
1. 秘密情報の定義と範囲
2. 秘密保持義務の内容
3. 秘密情報の使用目的
4. 秘密保持期間
5. 秘密情報の返還・廃棄
6. 例外事項（公知情報等）
7. 損害賠償
8. 準拠法と合意管轄

【出力形式】
- Markdown形式で出力（## 第X条 条項名、本文）
- 条項番号は「第1条」「第2条」...と順番に付ける
- 正式な法律文書の文体（甲・乙を使用）
- AIによる断定的な法的判断は記載しない
""",
    "sale": """\
あなたは日本の契約実務に精通した法律文書作成の専門家です。
以下の条件で売買契約書を作成してください。

【必ず含める条項】
1. 売買の目的物
2. 売買代金と支払方法
3. 引渡条件
4. 所有権移転
5. 危険負担
6. 瑕疵担保責任（契約不適合責任）
7. 解除条件
8. 準拠法と合意管轄

【出力形式】
- Markdown形式で出力（## 第X条 条項名、本文）
- 条項番号は「第1条」「第2条」...と順番に付ける
- 正式な法律文書の文体（売主・買主または甲・乙を使用）
- AIによる断定的な法的判断は記載しない
""",
    "tos": """\
あなたは日本のWebサービス運営実務に精通した法律文書作成の専門家です。
以下の条件で利用規約を作成してください。

【必ず含める条項】
1. サービスの目的と提供内容
2. 利用資格と登録
3. 禁止事項
4. 免責事項
5. 知的財産権
6. 個人情報の取扱い
7. サービスの変更・停止
8. 準拠法と管轄

【出力形式】
- Markdown形式で出力（## 第X条 条項名、本文）
- 条項番号は「第1条」「第2条」...と順番に付ける
- 利用者とサービス提供者の関係を明確にした文体
- AIによる断定的な法的判断は記載しない
""",
    "privacy": """\
あなたは日本の個人情報保護法に精通した法律文書作成の専門家です。
以下の条件でプライバシーポリシーを作成してください。

【必ず含める項目】
1. 取得する個人情報の項目
2. 個人情報の利用目的
3. 個人情報の第三者提供
4. 個人情報の管理
5. Cookieの使用について
6. 開示・訂正・削除の請求方法
7. 問い合わせ先
8. ポリシーの改定

【出力形式】
- Markdown形式で出力（## 項目名、本文）
- 個人情報保護法の要件を満たす内容
- わかりやすい文体（利用者目線）
- AIによる断定的な法的判断は記載しない
""",
    "employment": """\
あなたは日本の労働法・雇用実務に精通した法律文書作成の専門家です。
以下の条件で雇用契約書を作成してください。

【必ず含める条項】
1. 雇用期間（有期・無期）
2. 業務内容
3. 勤務場所
4. 勤務時間・休日
5. 賃金と支払方法
6. 社会保険・福利厚生
7. 解雇・退職の手続き
8. 準拠法と合意管轄

【出力形式】
- Markdown形式で出力（## 第X条 条項名、本文）
- 条項番号は「第1条」「第2条」...と順番に付ける
- 労働基準法に準拠した記載であることを意識した文体（使用者・労働者を使用）
- AIによる断定的な法的判断は記載しない
""",
    "rent": """\
あなたは日本の不動産・賃貸借契約実務に精通した法律文書作成の専門家です。
以下の条件で賃貸借契約書を作成してください。

【必ず含める条項】
1. 賃貸物件の表示
2. 賃料と支払方法
3. 敷金・保証金
4. 契約期間
5. 使用目的
6. 修繕・費用負担
7. 禁止事項
8. 解約・明渡条件
9. 準拠法と合意管轄

【出力形式】
- Markdown形式で出力（## 第X条 条項名、本文）
- 条項番号は「第1条」「第2条」...と順番に付ける
- 正式な法律文書の文体（貸主・借主または甲・乙を使用）
- AIによる断定的な法的判断は記載しない
""",
}


def _build_user_message(doc_type: str, params: dict) -> str:
    """文書種別と入力パラメータからユーザーメッセージを組み立てる"""
    doc_info = DOC_TYPES.get(doc_type, {})
    doc_name = doc_info.get("name", doc_type)

    lines = [f"以下の情報をもとに、{doc_name}を作成してください。", "", "【共通情報】"]

    party_a = params.get("party_a_name", "甲")
    party_b = params.get("party_b_name", "乙")
    effective_date = params.get("effective_date", "")
    governing_law = params.get("governing_law", "日本法")
    jurisdiction = params.get("jurisdiction", "東京地方裁判所")

    lines.append(f"- 甲（第一当事者）: {party_a}")
    lines.append(f"- 乙（第二当事者）: {party_b}")
    if effective_date:
        lines.append(f"- 契約日: {effective_date}")
    lines.append(f"- 準拠法: {governing_law}")
    lines.append(f"- 合意管轄: {jurisdiction}")

    lines.append("")
    lines.append("【種別固有情報】")

    if doc_type == "commission":
        lines.append(f"- 業務内容: {params.get('work_content', '（記載なし）')}")
        lines.append(f"- 報酬: {params.get('compensation', '（記載なし）')}")
        lines.append(f"- 期間: {params.get('period', '（記載なし）')}")
        lines.append(f"- 知財帰属: {params.get('ip_ownership', '受託者に帰属')}")

    elif doc_type == "nda":
        lines.append(f"- 目的: {params.get('purpose', '（記載なし）')}")
        lines.append(f"- 対象情報の範囲: {params.get('info_scope', '（記載なし）')}")
        lines.append(f"- 期間: {params.get('period', '（記載なし）')}")

    elif doc_type == "sale":
        lines.append(f"- 商品名: {params.get('product_name', '（記載なし）')}")
        lines.append(f"- 金額: {params.get('amount', '（記載なし）')}")
        lines.append(f"- 引渡条件: {params.get('delivery_conditions', '（記載なし）')}")

    elif doc_type == "tos":
        lines.append(f"- サービス名: {params.get('service_name', '（記載なし）')}")
        lines.append(f"- 禁止事項: {params.get('prohibited_items', '（記載なし）')}")
        lines.append(f"- 免責事項: {params.get('disclaimer', '（記載なし）')}")

    elif doc_type == "privacy":
        lines.append(f"- 取得する個人情報の項目: {params.get('collected_info', '（記載なし）')}")
        lines.append(f"- 利用目的: {params.get('usage_purpose', '（記載なし）')}")
        lines.append(f"- 問い合わせ先: {params.get('contact_info', '（記載なし）')}")

    elif doc_type == "employment":
        lines.append(f"- 業務内容: {params.get('work_content', '（記載なし）')}")
        lines.append(f"- 給与: {params.get('salary', '（記載なし）')}")
        lines.append(f"- 勤務地: {params.get('work_location', '（記載なし）')}")
        lines.append(f"- 期間: {params.get('period', '（記載なし）')}")

    elif doc_type == "rent":
        lines.append(f"- 物件: {params.get('property', '（記載なし）')}")
        lines.append(f"- 賃料: {params.get('rent', '（記載なし）')}")
        lines.append(f"- 期間: {params.get('period', '（記載なし）')}")
        lines.append(f"- 敷金: {params.get('deposit', '（記載なし）')}")

    return "\n".join(lines)


def ensure_disclaimer(text: str) -> str:
    """免責文言が冒頭になければ挿入して返す"""
    if DISCLAIMER not in text:
        return f"{DISCLAIMER}\n\n{text}"
    return text


async def generate_legal_document(doc_type: str, params: dict) -> str:
    """GPT-4o-miniで法的文書を生成する"""
    system_prompt = SYSTEM_PROMPTS.get(doc_type, SYSTEM_PROMPTS["commission"])
    user_message = _build_user_message(doc_type, params)

    client = _get_client()
    resp = await client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
        temperature=0.3,
        max_tokens=3000,
    )
    generated = resp.choices[0].message.content or ""
    return ensure_disclaimer(generated)


def build_docx(text: str) -> bytes:
    """MarkdownテキストからWordファイルを生成してbytesで返す"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docxが未インストールです。pip install python-docx でインストールしてください。")

    doc = Document()

    # スタイル設定
    style = doc.styles["Normal"]
    style.font.name = "MS Gothic"
    style.font.size = Pt(10.5)

    lines = text.split("\n")
    for line in lines:
        line = line.rstrip()

        if line.startswith("## "):
            # 見出し2レベル（条項タイトル）
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("# "):
            # 見出し1レベル（文書タイトル）
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("### "):
            heading = doc.add_heading(line[4:], level=3)
        elif line.startswith("※ "):
            # 免責文言を強調表示
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
            run.bold = True
        elif line.strip():
            # 通常段落
            doc.add_paragraph(line)
        else:
            # 空行
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_pdf(text: str) -> bytes:
    """MarkdownテキストからPDFを生成してbytesで返す"""
    try:
        from weasyprint import HTML as WeasyHTML
    except ImportError:
        raise ImportError(
            "WeasyPrintが未インストールです。pip install WeasyPrint でインストールしてください。"
        )

    # MarkdownをシンプルなHTMLに変換
    html_body = _markdown_to_html(text)

    html_content = f"""<!DOCTYPE html>
<html lang="ja">
<head>
<meta charset="utf-8">
<style>
  @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+JP:wght@400;700&display=swap');
  body {{
    font-family: 'Noto Sans JP', 'Hiragino Kaku Gothic ProN', 'Yu Gothic', sans-serif;
    font-size: 10.5pt;
    line-height: 1.8;
    margin: 2.5cm;
    color: #333;
  }}
  h1 {{ font-size: 14pt; text-align: center; margin-bottom: 1em; }}
  h2 {{ font-size: 11pt; margin-top: 1.5em; margin-bottom: 0.5em; border-bottom: 1px solid #ccc; }}
  h3 {{ font-size: 10.5pt; margin-top: 1em; }}
  p {{ margin: 0.5em 0; text-indent: 1em; }}
  .disclaimer {{
    background: #fff8e1;
    border-left: 4px solid #f59e0b;
    padding: 0.8em 1em;
    margin-bottom: 1.5em;
    font-weight: bold;
    color: #92400e;
    text-indent: 0;
  }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""

    pdf_bytes = WeasyHTML(string=html_content).write_pdf()
    return pdf_bytes


def _markdown_to_html(text: str) -> str:
    """シンプルなMarkdown→HTML変換"""
    lines = text.split("\n")
    html_lines = []
    in_paragraph = False

    for line in lines:
        line_stripped = line.strip()

        if line_stripped.startswith("# "):
            if in_paragraph:
                html_lines.append("</p>")
                in_paragraph = False
            content = line_stripped[2:]
            html_lines.append(f"<h1>{_escape_html(content)}</h1>")

        elif line_stripped.startswith("## "):
            if in_paragraph:
                html_lines.append("</p>")
                in_paragraph = False
            content = line_stripped[3:]
            html_lines.append(f"<h2>{_escape_html(content)}</h2>")

        elif line_stripped.startswith("### "):
            if in_paragraph:
                html_lines.append("</p>")
                in_paragraph = False
            content = line_stripped[4:]
            html_lines.append(f"<h3>{_escape_html(content)}</h3>")

        elif line_stripped.startswith("※ "):
            if in_paragraph:
                html_lines.append("</p>")
                in_paragraph = False
            html_lines.append(f'<p class="disclaimer">{_escape_html(line_stripped)}</p>')

        elif line_stripped == "":
            if in_paragraph:
                html_lines.append("</p>")
                in_paragraph = False

        else:
            if not in_paragraph:
                html_lines.append("<p>")
                in_paragraph = True
            else:
                html_lines.append("<br>")
            html_lines.append(_escape_html(line_stripped))

    if in_paragraph:
        html_lines.append("</p>")

    return "\n".join(html_lines)


def _escape_html(text: str) -> str:
    """HTMLエスケープ"""
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
    )
